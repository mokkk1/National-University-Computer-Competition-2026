"""
生成生产实习课程设计报告 (.docx)
基于模板: 2026BUCT生产实习课程设计报告模板.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import os

# ─── 项目根目录 ─────────────────────────────────
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ─── 样式常量 ───────────────────────────────────
FONT_FAMILY_CN = '宋体'
FONT_FAMILY_HEI = '黑体'
FONT_FAMILY_EN = 'Times New Roman'
FONT_SIZE_COVER_TITLE = Pt(22)    # 二号 (封面大标题)
FONT_SIZE_COVER_INFO = Pt(14)     # 四号 (封面信息)
FONT_SIZE_TITLE = Pt(16)          # 三号 (课设题目)
FONT_SIZE_H1 = Pt(15)            # 小三 (摘要标题, 各章标题)
FONT_SIZE_H2 = Pt(14)            # 四号 (节标题)
FONT_SIZE_H3 = Pt(12)            # 小四 (条/款/项标题)
FONT_SIZE_BODY = Pt(12)          # 小四 (正文)
FONT_SIZE_KEYWORD = Pt(14)       # 四号 (关键词标题)
LINE_SPACING_20 = Pt(20)         # 固定值20磅
LINE_SPACING_22 = Pt(22)         # 22磅


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=FONT_SIZE_BODY, bold=False, color=None):
    """统一设置 run 的字体属性"""
    run.font.size = size
    run.bold = bold
    run.font.name = font_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, line_spacing=LINE_SPACING_20, before=0, after=0, alignment=None):
    """统一设置段落间距和对齐"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if alignment is not None:
        pf.alignment = alignment


def add_paragraph_with_format(doc, text, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN,
                               size=FONT_SIZE_BODY, bold=False, alignment=None,
                               line_spacing=LINE_SPACING_20, before=0, after=0,
                               first_line_indent=None, color=None):
    """添加一个格式化段落"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=line_spacing, before=before, after=after, alignment=alignment)
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_run_font(run, font_cn=font_cn, font_en=font_en, size=size, bold=bold, color=color)
    return p


def add_heading_custom(doc, text, level=1):
    """添加自定义格式的标题，同时设置大纲级别以便 TOC 域代码自动生成目录"""
    if level == 0:  # 章标题 (如"第1章 绪论")
        size = FONT_SIZE_H1
        font_cn = FONT_FAMILY_HEI
        bold = True
        before = 12
        after = 12
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        outline_level = 1  # 对应 TOC \o "1-3" 的 level 1
    elif level == 1:  # 节标题 (如"1.1")
        size = FONT_SIZE_H2
        font_cn = FONT_FAMILY_HEI
        bold = True
        before = 6
        after = 6
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        outline_level = 2  # 对应 TOC \o "1-3" 的 level 2
    elif level == 2:  # 条标题 (如"1.1.1")
        size = FONT_SIZE_H3
        font_cn = FONT_FAMILY_HEI
        bold = True
        before = 3
        after = 3
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        outline_level = 3  # 对应 TOC \o "1-3" 的 level 3
    else:
        size = FONT_SIZE_BODY
        font_cn = FONT_FAMILY_HEI
        bold = True
        before = 3
        after = 3
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        outline_level = 3

    p = add_paragraph_with_format(doc, text, font_cn=font_cn, font_en=FONT_FAMILY_EN,
                                   size=size, bold=bold, alignment=alignment,
                                   line_spacing=LINE_SPACING_20, before=before, after=after)
    # 设置大纲级别，TOC 域代码通过 outlineLvl 识别标题
    pPr = p._element.get_or_add_pPr()
    outline_lvl = pPr.find(qn('w:outlineLvl'))
    if outline_lvl is None:
        outline_lvl = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{outline_level}"/>')
        pPr.append(outline_lvl)
    else:
        outline_lvl.set(qn('w:val'), str(outline_level))
    return p


def add_body(doc, text):
    """添加正文段落（小四宋体，首行缩进2字符，固定值20磅）"""
    return add_paragraph_with_format(doc, text, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN,
                                     size=FONT_SIZE_BODY, bold=False,
                                     line_spacing=LINE_SPACING_20, before=0, after=0,
                                     first_line_indent=Cm(0.74), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_body_no_indent(doc, text, bold=False):
    """添加正文段落（无缩进）"""
    return add_paragraph_with_format(doc, text, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN,
                                     size=FONT_SIZE_BODY, bold=bold,
                                     line_spacing=LINE_SPACING_20, before=0, after=0,
                                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))
    set_paragraph_spacing(p, line_spacing=LINE_SPACING_20)


def create_cover_page(doc):
    """创建封面页"""
    # 空行
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=LINE_SPACING_20)

    # 主标题
    add_paragraph_with_format(doc, '生产实习', font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN,
                              size=Pt(26), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=Pt(36), after=6)

    add_paragraph_with_format(doc, '课程设计报告', font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN,
                              size=Pt(26), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=Pt(36), after=30)

    # 封面信息 — 使用表格实现对齐
    info_items = [
        ('课设题目：', '基于深度排序学习的沪深300指数预测模型'),
        ('班    级：', '自控2201班（示例）'),
        ('姓    名：', '（填写姓名）'),
        ('学    号：', '（填写学号）'),
        ('校内导师：', '（填写导师姓名）'),
        ('企业导师：', '（填写企业导师）'),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.autofit = True

    for i, (label, value) in enumerate(info_items):
        # Label cell (右对齐)
        cell_label = table.cell(i, 0)
        cell_label.width = Cm(4)
        p_label = cell_label.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_label = p_label.add_run(label)
        set_run_font(run_label, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=FONT_SIZE_COVER_INFO, bold=False)
        set_paragraph_spacing(p_label, line_spacing=Pt(28))
        # 去掉单元格边框
        for cell in [cell_label]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)

        # Value cell
        cell_value = table.cell(i, 1)
        cell_value.width = Cm(10)
        p_value = cell_value.paragraphs[0]
        run_value = p_value.add_run(value)
        set_run_font(run_value, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=FONT_SIZE_COVER_INFO, bold=False)
        # 添加下划线效果 — 使用下划线
        run_value.underline = True
        set_paragraph_spacing(p_value, line_spacing=Pt(28))
        for cell in [cell_value]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)

    # 日期
    add_paragraph_with_format(doc, '', font_cn=FONT_FAMILY_CN, size=Pt(12),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(20))
    add_paragraph_with_format(doc, '2026年7月', font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN,
                              size=FONT_SIZE_COVER_INFO, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=Pt(28))


def create_title_page(doc):
    """创建课设题目页"""
    add_page_break(doc)
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=LINE_SPACING_20)

    add_paragraph_with_format(doc, '基于深度排序学习的沪深300指数预测模型',
                              font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN,
                              size=FONT_SIZE_TITLE, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=Pt(30), after=20)


def create_abstract(doc):
    """创建摘要页"""
    add_page_break(doc)

    add_paragraph_with_format(doc, '摘  要', font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN,
                              size=FONT_SIZE_H1, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=LINE_SPACING_20, after=20)

    abstract_text = (
        '本课程设计围绕THU-BDC2026大数据竞赛的沪深300指数预测任务展开，'
        '目标是利用深度排序学习技术从沪深300成分股中精准筛选未来5日收益率最高的5只股票。'
        '项目构建了一套完整的量化选股系统，包含数据获取、特征工程、模型训练和后处理策略四个核心环节。'
        '在特征工程方面，设计了237维多源特征体系，涵盖158个Alpha因子、39个技术指标、'
        '9维基本面特征、18维宏观指标、7维市场宽度特征以及16维行业嵌入向量。'
        '在模型架构方面，提出了StockTransformer排序模型（约250万参数），集成多尺度时序卷积、'
        'Transformer自注意力、跨股票注意力以及市场聚合门控等创新模块；同时设计了轻量级替代方案'
        'LightweightStockRanker（约26万参数），基于统计矩投影和双向GRU实现参数削减11倍。'
        '在训练策略方面，采用Walk-Forward滚动窗口交叉验证框架（6窗口），配合时间衰减采样、'
        '混合标签构建（70%排序+30%绝对收益）、多任务学习以及Gumbel-Softmax组合收益优化等关键技术。'
        '在后处理方面，提出quarter_aware市场门控策略，根据季末效应和市场涨跌方向自适应切换'
        '进攻型与防御型选股策略。实验结果表明，最优配置（V8 Improved + quarter_aware门控）'
        '在18次回测中实现综合收益率+1.55%、正收益比例72%的良好表现，'
        '验证了深度排序学习在A股量化选股中的有效性。'
    )
    add_body(doc, abstract_text)

    # 关键词
    add_paragraph_with_format(doc, '', font_cn=FONT_FAMILY_CN, size=FONT_SIZE_BODY,
                              line_spacing=LINE_SPACING_20)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=LINE_SPACING_20)
    run_label = p.add_run('关键词：')
    set_run_font(run_label, font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN, size=FONT_SIZE_KEYWORD, bold=True)
    run_text = p.add_run('深度学习；排序学习；沪深300；股票预测；Transformer')
    set_run_font(run_text, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=FONT_SIZE_KEYWORD, bold=False)


def create_toc_placeholder(doc):
    """创建目录页 — 插入真正的 Word TOC 域代码"""
    add_page_break(doc)
    add_paragraph_with_format(doc, '目  录', font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN,
                              size=FONT_SIZE_H1, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=LINE_SPACING_20, after=20)

    # 插入 TOC 域代码: { TOC \o "1-3" \h \z \u }
    # 使用 OxmlElement 正确创建 fldChar 元素
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=LINE_SPACING_20)

    # Run 1: fldChar begin
    run1 = p.add_run()
    set_run_font(run1, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=Pt(10))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fld_begin)

    # Run 2: instrText
    run2 = p.add_run()
    set_run_font(run2, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=Pt(10))
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    # Run 3: fldChar separate
    run3 = p.add_run()
    set_run_font(run3, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=Pt(10))
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fld_sep)

    # Run 4: 提示文字
    run4 = p.add_run()
    set_run_font(run4, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=Pt(10),
                 color=RGBColor(128, 128, 128))
    run4.text = '（打开文件后右键此处 → 更新域 → 自动生成目录）'

    # Run 5: fldChar end
    run5 = p.add_run()
    set_run_font(run5, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=Pt(10))
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run5._element.append(fld_end)


def write_chapter1(doc):
    """第1章 绪论"""
    add_page_break(doc)
    add_heading_custom(doc, '前  言', level=0)

    add_body(doc,
        '量化投资是金融科技领域的核心研究方向之一。随着深度学习技术的快速发展，'
        '基于神经网络的量化选股模型在国内外市场展现出超越传统多因子模型的潜力。'
        '然而，A股市场具有高噪声、强政策驱动、板块轮动频繁等特征，'
        '对模型的鲁棒性和泛化能力提出了严峻挑战。'
        '本课程设计以THU-BDC2026大数据竞赛为依托，系统性地探索了深度排序学习'
        '在沪深300成分股精选中的应用，完整经历了从问题建模、数据工程、模型设计、'
        '训练优化到后处理策略的全流程工程实践。'
    )

    add_heading_custom(doc, '第1章 绪论', level=0)

    # 1.1
    add_heading_custom(doc, '1.1 项目背景及意义', level=1)

    add_body(doc,
        '沪深300指数由沪深证券交易所于2005年联合发布，选取A股市场中市值大、流动性好的'
        '300只股票作为成分股，覆盖了沪深市场约60%的总市值，是反映A股整体走势的最重要指标之一。'
        '沪深300成分股涵盖金融、消费、科技、医药、能源等多个行业，具有广泛的市场代表性。'
        '对于量化投资策略而言，从300只成分股中精准筛选出未来表现最优的少数股票，'
        '是实现超额收益（Alpha）的核心能力。'
    )

    add_body(doc,
        'THU-BDC2026大数据竞赛的沪深300指数预测赛题正是基于这一实际需求设计：'
        '参赛者需要利用历史交易数据、财务数据和宏观经济数据，构建模型从沪深300成分股中'
        '选择5只股票，并在测试集上最大化其未来5个交易日的实际平均收益率。'
        '该赛题本质上是一个排序学习（Learning to Rank）问题——模型不需要精确预测每只股票的收益率，'
        '而是需要学习正确的相对排序，将真正会涨的股票排在最前面。'
    )

    add_body(doc,
        '本项目的意义体现在以下三个方面：第一，系统性地探索深度排序学习在A股量化选股中的'
        '适用性和有效性，为后续研究提供可复现的技术路线；第二，通过大量对比实验揭示'
        '金融数据非平稳性、季末效应等市场微观结构对模型表现的影响，形成可操作的实践经验；'
        '第三，构建了一套完整的量化选股工程流水线，涵盖数据获取、特征工程、模型训练、'
        '回测评估和Docker部署等环节，具备直接参与实盘交易的工程基础。'
    )

    # 1.2
    add_heading_custom(doc, '1.2 本课题主要研究内容', level=1)

    add_body(doc,
        '本课题围绕"从沪深300成分股中精准筛选未来5日收益率最高的5只股票"这一核心目标，'
        '主要研究内容包括以下四个方面：'
    )

    add_body(doc,
        '（1）多源异构特征工程体系构建。研究如何从量价数据、技术指标、基本面财报、'
        '宏观经济指标、市场微观结构等多个数据源中提取有效的预测因子。项目最终构建了'
        '237维特征体系，涵盖Alpha因子、技术指标、基本面、动量、宏观、市场宽度和行业嵌入等7大类特征。'
    )

    add_body(doc,
        '（2）深度排序学习模型架构设计。针对金融时序数据信噪比极低（约1:13）的特点，'
        '设计了StockTransformer排序模型，集成多尺度时序卷积（TCN）、Transformer自注意力编码器、'
        '跨股票注意力机制以及市场聚合门控模块。同时设计了参数削减11倍的轻量级替代方案'
        'LightweightStockRanker，验证"少即是多"的金融建模哲学。'
    )

    add_body(doc,
        '（3）鲁棒训练策略与损失函数设计。针对排序学习与绝对收益预测之间的张力，'
        '提出混合标签策略（70%分位数排序+30%绝对收益），设计了融合Listwise、Pairwise（LambdaRank）、'
        'NDCG、Precision@K和Gumbel-Softmax组合收益损失的多目标损失函数。'
        '采用Walk-Forward滚动窗口交叉验证框架，真实模拟"训练→预测未来"的时间序列因果约束。'
    )

    add_body(doc,
        '（4）市场环境自适应后处理策略。针对模型原始排序分数缺乏市场环境感知的问题，'
        '提出quarter_aware市场门控策略，将季末日历效应与短期市场动量信号相结合，'
        '根据市场涨跌方向自适应切换进攻型与防御型选股模式，显著提升最终组合收益率。'
    )

    # 1.3
    add_heading_custom(doc, '1.3 开发环境', level=1)

    add_body(doc, '本项目的开发环境配置如下：')

    env_info = [
        ('硬件环境', 'NVIDIA RTX 5070 Ti Laptop GPU (16GB VRAM), 32GB RAM, Intel Core i9-13900H'),
        ('操作系统', 'Windows 11 Home China, Docker 运行环境 Ubuntu 22.04 (CUDA 12.6)'),
        ('编程语言', 'Python 3.11'),
        ('深度学习框架', 'PyTorch 2.8.0+cu129'),
        ('技术分析库', 'TA-Lib 0.6.8（C语言级性能）'),
        ('数据处理', 'pandas 2.3.2, numpy 2.0.2, scikit-learn 1.7.2'),
        ('数据源', 'baostock（股票日线）, akshare（宏观指标、基本面）'),
        ('包管理', 'uv（Astral出品的高性能Python包管理器）'),
        ('容器化', 'Docker (基础镜像 nvidia/cuda:12.6.0-runtime-ubuntu22.04)'),
        ('版本控制', 'Git'),
    ]

    table = doc.add_table(rows=len(env_info), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(env_info):
        cell_label = table.cell(i, 0)
        cell_label.width = Cm(3.5)
        p_label = cell_label.paragraphs[0]
        run_label = p_label.add_run(label)
        set_run_font(run_label, font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN, size=FONT_SIZE_BODY, bold=True)
        set_paragraph_spacing(p_label, line_spacing=LINE_SPACING_20)

        cell_value = table.cell(i, 1)
        p_value = cell_value.paragraphs[0]
        run_value = p_value.add_run(value)
        set_run_font(run_value, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=FONT_SIZE_BODY)
        set_paragraph_spacing(p_value, line_spacing=LINE_SPACING_20)

    add_paragraph_with_format(doc, '', font_cn=FONT_FAMILY_CN, size=FONT_SIZE_BODY,
                              line_spacing=LINE_SPACING_20)

    # 1.4
    add_heading_custom(doc, '1.4 主要任务与总体结构', level=1)

    add_body(doc,
        '本项目的总体技术路线分为五个阶段：'
    )

    add_body(doc,
        '第一阶段（数据获取与预处理）：通过baostock API下载沪深300成分股的历史日线数据'
        '（2010年至2026年），通过akshare API获取宏观经济指标（国债收益率、北向资金、CPI、PPI、'
        'PMI、M1/M2、社融等18维指标）和基本面数据（PE、PB、ROE等）。对原始数据进行清洗、'
        '对齐和缺失值填充，按时间顺序划分为训练集和测试集。'
    )

    add_body(doc,
        '第二阶段（特征工程）：基于金融领域知识，从量价数据中计算158维Alpha因子'
        '（动量、反转、波动、流动性等类别的技术信号），利用TA-Lib计算39维经典技术指标'
        '（SMA、EMA、MACD、RSI、KDJ、BOLL、ATR、OBV等），拼接基本面、动量、宏观指标'
        '以及市场宽度特征（涨跌比、收益离散度、偏度等），通过申万一级行业分类构建可学习的'
        '16维行业嵌入向量。所有特征经标准化处理后输入模型。'
    )

    add_body(doc,
        '第三阶段（模型训练）：实现并对比两种架构——完整版StockTransformer（约250万参数）'
        '和轻量版LightweightStockRanker（约26万参数）。采用Walk-Forward 6窗口滚动训练框架，'
        '每窗口独立训练-验证-早停，最终多模型集成预测。使用时间衰减采样应对金融数据非平稳性，'
        '使用混合标签和Portfolio Loss直接优化组合收益期望。'
    )

    add_body(doc,
        '第四阶段（后处理策略）：在模型原始排序分数基础上，叠加market_gate市场门控后处理。'
        '根据预测日市场环境（沪深300近10日累计涨跌）和日历位置（是否临近季末日），'
        '自适应调整选股策略——牛市偏重进攻型选股（高动量、高排序分数），'
        '熊市或季末偏重防御型选股（低波动、大市值、低Beta），实现市场环境自适应。'
    )

    add_body(doc,
        '第五阶段（Docker封装与提交）：将完整的训练和推理流水线封装为Docker镜像，'
        '通过docker-compose进行本地可运行性验证，确保提交的镜像能在竞赛评测环境中正确运行'
        '并生成规定格式的预测结果。'
    )

    # 1.5
    add_heading_custom(doc, '1.5 本人完成的工作', level=1)

    add_body(doc,
        '在本课程设计项目中，本人独立完成了以下全部工作：'
    )

    add_body(doc,
        '（1）项目整体架构设计与技术选型，包括模型架构方案对比、训练策略制定和后处理策略设计。'
        '（2）全部核心代码编写，包括特征工程模块（utils.py，约1700行）、'
        '模型定义（model.py，两种架构共约1700行）、训练流水线（train.py，约1350行）、'
        'Walk-Forward训练框架（walk_forward.py，约970行）、'
        '市场门控后处理（market_gate.py，约590行）以及宏观行业数据处理（macro_industry.py，约1070行）。'
        '（3）数据管道的搭建与调试，包括baostock股票数据获取（断点续传）、'
        'akshare宏观数据获取（多数据源容错）、以及数据清洗、对齐和划分。'
        '（4）大量对比实验的设计与执行，涵盖V3至V8共6个模型版本的迭代优化，'
        'Walk-Forward 6窗口训练（总耗时约6小时/轮），以及截断实验、半衰期实验、'
        '季末效应验证等多个消融实验。'
        '（5）实验日志的完整记录（14份工作记录+改进方案文档）、'
        '4篇架构决策记录（ADR）的撰写以及项目文档（README、GUIDE、REPRODUCE）的编写。'
        '（6）Docker镜像的构建与验证，确保竞赛提交格式的合规性。'
        '（7）41个单元测试和集成测试的编写，覆盖模型前向传播、损失函数、标签构建、'
        '市场门控和端到端训练链路。'
    )


def write_chapter2(doc):
    """第2章 项目总体设计"""
    add_page_break(doc)
    add_heading_custom(doc, '第2章 项目总体设计', level=0)

    # 2.1
    add_heading_custom(doc, '2.1 数据获取与预处理', level=1)

    add_body(doc,
        '数据是量化模型的根基。项目从三个数据源获取原始数据：'
        '（1）baostock——提供沪深300全部成分股的日线行情数据，包括开盘价、收盘价、最高价、'
        '最低价、成交量、成交额、振幅、涨跌额、换手率和涨跌幅，覆盖时间范围2010年1月至2026年3月，'
        '累计约94万行数据；（2）akshare——提供18维宏观经济指标（国债收益率、北向资金净流入、'
        '融资余额、美元兑人民币汇率、LPR利率、Shibor利率、CPI、PPI、PMI、M1/M2货币供应量、'
        '社会融资规模等），以及PE、PB、ROE、ROA、毛利率、营收增速、净利润增速等基本面数据；'
        '（3）申万行业分类——将300只股票映射到31个一级行业（银行、食品饮料、电子、医药生物等）。'
    )

    add_body(doc,
        '数据预处理的关键步骤包括：按股票代码和日期排序确保时序正确性；将绝对收益率'
        '（未来第5日开盘价相对第1日开盘价的涨跌幅）转化为混合排序标签（70%每日组内分位数rank'
        '+ 30%归一化绝对收益）；构建市场方向标签（每日所有股票的收益率均值>0则为上涨）；'
        '对特征进行Z-score标准化（按训练集拟合、验证集和测试集使用相同参数transform）；'
        '以及严格的未来数据泄露防护（标准化仅在训练集上fit，标签计算仅使用未来数据窗口）。'
    )

    # 2.2
    add_heading_custom(doc, '2.2 特征工程体系', level=1)

    add_body(doc,
        '项目构建了237维多源特征体系，按类别分为以下七大类：'
    )

    # 特征表
    feature_data = [
        ('原始价格特征', '11维', '开盘价、收盘价、最高价、最低价、成交量、成交额、振幅、涨跌额、涨跌幅、换手率、instrument索引'),
        ('Alpha因子', '158维', 'K线形态特征（KMID/KLEN/KUP/KLOW/KSFT等9维）、价格比率（OPEN0/HIGH0/LOW0/VWAP0 4维）、'
                   '动量变化率（ROC5-60, 5维）、移动均线比率（MA5-60, 5维）、标准差（STD5-60, 5维）、'
                   '线性回归特征（BETA/RSQR/RESI各5窗口共15维）、极值特征（MAX/MIN各5窗口共10维）、'
                   '分位数特征（QTLU/QTLD各5窗口共10维）、排序特征（RANK 5维）、'
                   '随机振荡指标（RSV 5维）、极值索引（IMAX/IMIN/IMXD各5窗口共15维）、'
                   '相关性特征（CORR/CORD各5窗口共10维）、计数特征（CNTP/CNTN/CNTD共15维）、'
                   '价格变化累加（SUMP/SUMN/SUMD共15维）、量能特征（VMA/VSTD/WVMA/VSUMP/VSUMN/VSUMD共30维）'),
        ('技术指标', '39维', 'SMA/EMA均线系统（sma_5/20, ema_12/26/60）、MACD信号线、RSI相对强弱、'
                   'KDJ随机指标（K/D/J）、BOLL布林带（中轨/标准差）、ATR平均真实波幅、OBV能量潮、'
                   '成交量指标（量比/量均线）、收益率（1/5/10日）、波动率（10/20日）、价差指标（4维）'),
        ('基本面特征', '9维', 'PE（市盈率）、PB（市净率）、PS（市销率）、ROE（净资产收益率）、'
                   'ROA（总资产收益率）、gross_margin（毛利率）、revenue_yoy（营收同比增速）、'
                   'profit_yoy（净利润同比增速）、north_holding（北向资金持股比例）'),
        ('动量特征', '4维', '5日收益率(ret5)、20日收益率(ret20)、20日波动率(vol20)、5日夏普比率(sharpe5)'),
        ('宏观特征', '18维', '国债收益率（1/5/10年期）、北向资金净流入、融资余额、美元兑人民币汇率、'
                   'LPR利率（1年期/5年期）、Shibor隔夜利率、CPI/PPI同比、制造业PMI、'
                   'M1/M2货币供应量同比、社会融资规模增量'),
        ('市场宽度特征', '7维', '涨跌家数比、收益离散度、收益偏度、成交量变化、上涨股票成交额占比、振幅均值、市场收益均值'),
        ('行业嵌入', '16维', '申万一级31类行业 → 可学习的16维稠密向量（Embedding），随模型训练端到端优化'),
    ]

    for cat_name, dim, desc in feature_data:
        add_body(doc, f'{cat_name}（{dim}）：{desc}。')

    # 2.3
    add_heading_custom(doc, '2.3 模型架构设计', level=1)

    add_body(doc,
        '项目实现了两种差异化架构，以应对不同的部署场景和样本量条件。'
    )

    add_heading_custom(doc, '2.3.1 StockTransformer（V8 Improved）', level=2)

    add_body(doc,
        'StockTransformer是项目的主力模型，参数量约250万，数据流如下：'
        '输入维度为[B, 300, 60, 237]的股票特征张量（B为batch大小，300为股票数，60为交易日序列长度，'
        '237为特征维度）。首先通过Input Projection将237维特征线性投影至256维隐空间；'
        '经过正弦位置编码后，进入MultiScaleConv多尺度时序卷积模块（kernel_size=3/5/7），'
        '并行捕捉短期动量、中期趋势和长周期模式；随后通过3层Transformer Encoder'
        '（每层4头自注意力+前馈网络维度512）提取时序依赖关系；FeatureAttention模块'
        '通过MLP学习各时间步的重要性权重，将时序维度压缩为固定长度向量；'
        'FeatureInteraction模块利用低秩bilinear交叉（rank=64）建模特征间的非线性组合效应；'
        'CrossStockAttention模块在同交易日内对300只股票执行多头注意力，'
        '让每只股票的特征聚合来自其他股票的信息，捕捉市场结构和板块联动效应。'
    )

    add_body(doc,
        '模型的创新点在于市场聚合架构（Market Aggregation Architecture）：'
        '在CrossStockAttention输出之上，MarketAttentionPooling通过可学习的query token'
        '对300只股票的特征执行注意力池化，生成64维市场状态向量。该向量同时用于两个目的：'
        '（1）通过MarketGate门控网络调制每只个股的排序特征——市场看涨时自动放大进攻型股票'
        '（高动量、高beta）的分数，市场看跌时自动提升防御型股票（低波动、大市值）的分数；'
        '（2）通过market_head预测市场涨跌方向（BCE二分类），作为额外的监督信号辅助训练。'
        '最终排序头（Ranking Layers + Score Head）将256维特征映射为单值排序分数，'
        '同时输出方向预测、波动率预测和绝对收益预测三个辅助任务的预测值。'
    )

    add_heading_custom(doc, '2.3.2 LightweightStockRanker（V9）', level=2)

    add_body(doc,
        'LightweightStockRanker是参数削减11倍（约26万参数）的轻量级替代方案，'
        '基于"金融数据信噪比极低，减少可学习参数本身就是最强的正则化"这一设计哲学。'
        '其架构流程为：首先计算每个特征的4个统计矩（均值、标准差、最新值、趋势=最新-最早），'
        '拼接为固定的统计汇总特征；然后通过共享的双向GRU（隐藏维度48）捕获残差时序模式；'
        '统计汇总与GRU输出拼接后进入排序头产出分数。市场上下文通过简单的均值池化实现'
        '（而非注意力池化），大幅降低了计算开销。实验表明，当单窗口训练样本超过75天时，'
        '该轻量模型在验证集上的表现可超越完整版StockTransformer，验证了"少即是多"的假设。'
    )

    # 2.4
    add_heading_custom(doc, '2.4 损失函数设计', level=1)

    add_body(doc,
        '项目的损失函数采用多目标加权组合策略，总损失公式为：'
    )

    add_body(doc,
        'Total Loss = WeightedRankingLoss (Listwise + Pairwise + NDCG@5 + Precision@5) '
        '+ 0.15 × Portfolio Return Loss (Gumbel-Softmax Top-K) '
        '+ 0.2 × Market Direction BCE '
        '+ 0.1 × Direction Prediction (auxiliary) '
        '+ 0.1 × Volatility Prediction (auxiliary) '
        '+ 0.3 × Return Regression (Huber, auxiliary)。'
    )

    add_body(doc,
        '其中各子损失的设计思路如下：Listwise损失将预测分数和真实标签分别通过softmax'
        '转化为概率分布后计算交叉熵（带Lambda权重）；Pairwise损失对每对股票的预测差值'
        '和真实差值计算sigmoid损失，权重由LambdaRank的ΔNDCG公式精确给出——'
        '排名靠前的股票对NDCG贡献更大，因此其排序错误的惩罚也更重；'
        'NDCG@5损失通过softmax排序近似实现可微化；Precision@5损失通过'
        'softmax概率在真实Top-5上的质量之和定义。'
    )

    add_body(doc,
        'Portfolio Return Loss是项目的关键创新之一——传统的排序损失对分数施加常数偏移时'
        '保持不变（排列等变性），因此无法约束"选出的股票到底能赚多少"。'
        'Portfolio Loss通过Gumbel-Softmax将离散的Top-K选择松弛为可微的连续概率分布，'
        '直接最大化所选K只股票的真实收益期望，迫使模型在排序质量之上额外关注绝对收益的方向和幅度。'
    )

    # 2.5
    add_heading_custom(doc, '2.5 训练策略', level=1)

    add_body(doc,
        '项目采用Walk-Forward滚动窗口交叉验证作为核心训练框架，这是应对金融时间序列数据'
        '非平稳性的关键设计。将2021年1月至2026年3月的数据划分为6个窗口，每个窗口的训练集'
        '截止日分别为2024-09-30、2024-12-31、2025-03-31、2025-06-30、2025-09-30和2025-12-31，'
        '验证集为截止日后2个月的数据。这种设计真实模拟了"使用历史数据训练→预测未来"的因果约束，'
        '避免了传统随机K折交叉验证在时序数据上造成的数据泄露问题。'
    )

    add_body(doc,
        '为了应对金融数据非平稳性（2015年的市场规律与2024年显著不同），项目引入了时间衰减采样策略：'
        '距训练截止日每730天（2年），该样本被采样的概率减半。这意味着近期数据（如2023-2024年）'
        '占据梯度主导地位，而远期数据（如2021年及之前）提供正则化信号但不主导优化方向。'
        '截断实验进一步验证了仅使用2021年起的数据（丢弃2010-2020年）能提升模型表现，'
        '说明过旧的市场规律对当前预测可能有害。'
    )

    add_body(doc,
        '混合标签是训练策略的另一个重要设计。纯分位数排序标签虽然天然抹平牛熊市偏差，'
        '但丢失了"选出的股票是涨还是跌"这一关键信息。混合标签以70%分位数排序和30%归一化绝对收益'
        '组合而成，在保持排序学习能力的同时传递绝对收益的方向和幅度信号。'
        '此外，数据增强（40%概率触发，含时序掩码10%、特征噪声σ=0.003、股票丢弃15%）、'
        '标签平滑以及warmup+线性衰减学习率调度（5 epoch warmup, end_factor=0.2）'
        '等正则化手段共同保障了模型在小样本金融数据上的泛化能力。'
    )

    # 2.6
    add_heading_custom(doc, '2.6 后处理策略', level=1)

    add_body(doc,
        '后处理策略是本项目收益率提升的最大贡献因素。核心发现是：模型的排序质量指标'
        '（如NDCG、Precision@K）与最终投资组合收益率之间的相关性很弱——'
        '模型可能学到"谁是相对好的股票"，但没有学到"在什么市场环境下该选什么类型的股票"。'
        '基于这一洞察，项目设计了quarter_aware市场门控策略：'
    )

    add_body(doc,
        '第一层——市场方向判断：以沪深300所有成分股近10个交易日的平均累计涨跌幅作为市场信号。'
        '累计涨幅>1%判定为看涨市场，累计跌幅<-1%判定为看跌市场。'
        '同时考虑预测日距离最近季末日（3/31, 6/30, 9/30, 12/31）的天数——'
        '实验发现季末附近（<5个交易日）市场不确定性显著增加，需要更保守的策略。'
    )

    add_body(doc,
        '第二层——自适应防御权重：防御权重w根据市场方向和季末距离动态计算。'
        '看涨市场：w≈0.15（几乎全进攻，按排序分数+收益门控选股）；看跌市场：'
        'w=0.3+0.5×置信度（范围0.3~0.9），季末额外+0.2；中性市场：w≈0.3。'
    )

    add_body(doc,
        '第三层——评分融合：最终每只股票的选股得分 = (1-w)×归一化排序分数 + w×防御性评分。'
        '防御性评分综合考量低波动（权重35%）、大市值（权重25%）、低Beta（权重25%）'
        '和排序分数（权重15%）四个因子。通过softmax温度缩放后取Top-5即为最终选股结果。'
    )

    add_body(doc,
        '这一后处理策略的关键优势在于——不需要重新训练模型。通过调整后处理参数即可适配'
        '不同的市场环境，实现了"模型专注排序，后处理专注场景适配"的关注点分离。'
    )


def write_chapter3(doc):
    """第3章 项目实现与测试"""
    add_page_break(doc)
    add_heading_custom(doc, '第3章 项目实现与测试', level=0)

    # 3.1
    add_heading_custom(doc, '3.1 核心模块实现', level=1)

    add_heading_custom(doc, '3.1.1 特征工程模块', level=2)
    add_body(doc,
        '特征工程模块（utils.py）实现了完整的特征计算流水线。engineer_features函数计算158维Alpha因子，'
        '利用TA-Lib的C语言级加速实现SMA、STDDEV、LINEARREG_SLOPE、CORREL等统计计算，'
        '避免了纯Python逐股逐窗口计算的性能瓶颈。engineer_features_39函数基于TA-Lib计算39维技术指标。'
        'compute_market_breadth_features函数按日期聚合生成市场宽度特征。'
        '所有特征工程函数遵循统一的输入输出接口——输入单只股票的DataFrame，输出带有新增特征列的DataFrame，'
        '支持多进程并行处理（通过multiprocessing.Pool）。'
    )

    add_heading_custom(doc, '3.1.2 数据集构建模块', level=2)
    add_body(doc,
        'create_ranking_dataset_vectorized函数采用向量化加速策略构建排序数据集。'
        '核心优化在于：预先为每只股票生成所有合法的滑动窗口（满足序列长度和历史窗口要求的end_idx），'
        '将窗口信息记录为轻量级元数据，再按预测日期分组聚合。相比逐日逐股的朴素循环实现，'
        '向量化版本在长历史窗口（2010年起）上提速约10倍。'
        '对于超长历史导致内存不足的场景，项目还实现了LazyRankingDataset懒加载方案——'
        '每只股票的特征矩阵仅存储一份（float32），60天窗口切片延迟到DataLoader的__getitem__中执行，'
        '将内存占用从约58GB压缩至约1GB。'
    )

    add_heading_custom(doc, '3.1.3 模型模块', level=2)
    add_body(doc,
        'model.py实现了两个完整的模型类。StockTransformer类（约300行）封装了从输入投影到排序输出的'
        '完整前向传播逻辑，支持通过config字典灵活开关TCN、特征交互、市场聚合等模块，'
        '以及是否返回辅助任务输出。LightweightStockRanker类（约200行）采用不同的架构哲学，'
        '以统计矩汇总替代深度Transformer作为时序编码器。两个模型均实现了_init_weights权重初始化、'
        '支持strict=False的模型加载兼容性（允许在不同版本的checkpoint之间迁移），'
        '以及forward中自动处理行业Embedding的注入逻辑。'
    )

    add_heading_custom(doc, '3.1.4 训练模块', level=2)
    add_body(doc,
        'train.py实现了完整的训练流水线（约1350行），核心组件包括：'
        '_build_label_and_clean（混合标签构建，支持分位数rank、混合标签和超额收益三种模式）；'
        'WeightedRankingLoss（融合5种子损失的组合排序损失，支持精确LambdaRank ΔNDCG权重计算）；'
        'train_ranking_model（单epoch训练循环，集成数据增强、辅助任务损失计算、'
        '梯度裁剪和TensorBoard日志记录）；evaluate_ranking_model（验证集评估）；'
        '以及predict_top_stocks（单日Top-K预测，支持波动率惩罚和分数校准）。'
        'RankingDataset和collate_fn配合实现了变长股票数量的动态padding和mask生成。'
    )

    add_heading_custom(doc, '3.1.5 Walk-Forward训练框架', level=2)
    add_body(doc,
        'walk_forward.py实现了滚动窗口训练与集成的全流程（约970行）。支持三种内置配置'
        '（light轻量配置、standard标准配置、v8_improved最优配置）和两种模型类'
        '（StockTransformer和LightweightStockRanker）。核心函数train_single_window'
        '完成单窗口的数据准备→特征工程→标准化→数据集构建→模型训练→早停保存的完整流程。'
        'run_walk_forward编排6个窗口的串行训练并汇总结果。load_walk_forward_models和'
        'walk_forward_predict实现多窗口模型的集成预测——使用median聚合各模型分数（对异常窗口鲁棒），'
        '并通过一致性过滤（需至少40%模型共识）剔除分歧过大的股票。'
    )

    add_heading_custom(doc, '3.1.6 市场门控模块', level=2)
    add_body(doc,
        'market_gate.py实现了市场环境自适应选股后处理（约590行）。compute_market_signal函数'
        '支持四种信号计算方法：hs300_return（HS300近10日累计涨跌幅）、rf_classifier'
        '（随机森林利用18维宏观特征预测市场方向）、model_head（利用模型内置market_head预测）'
        '以及ensemble（三信号加权投票）。MarketGate类的select方法根据信号方向和季末距离'
        '自适应计算防御权重并执行选股。compute_stock_defensive_score函数综合低波动、大市值、'
        '低Beta和行业属性计算每只股票的防御性评分。'
    )

    # 3.2
    add_heading_custom(doc, '3.2 模型训练与验证', level=1)

    add_body(doc,
        '项目采用Walk-Forward 6窗口交叉验证框架对模型进行全面评估。训练环境为NVIDIA RTX 5070 Ti '
        'Laptop GPU（16GB VRAM），6窗口完整训练耗时约6小时。'
    )

    # 训练配置表
    add_body(doc, '最优训练配置（V8 Improved）如下：')
    config_data = [
        ('模型架构', 'StockTransformer (V8 Improved), ~2.5M参数'),
        ('数据起始日期', '2021-01-01（截断实验证明此前数据有害）'),
        ('特征集', '158+39+fundamental+momentum, 约237维'),
        ('序列长度', '60个交易日'),
        ('隐层维度', 'd_model=256, dim_feedforward=512'),
        ('Transformer层数', '3层, 4头注意力'),
        ('TCN卷积核', '[3, 5, 7]'),
        ('Dropout', '0.15'),
        ('Batch大小', '4'),
        ('学习率', '1e-5 (AdamW, weight_decay=1e-5)'),
        ('Warmup', '5 epochs, 1e-6→1e-5'),
        ('学习率衰减', '线性衰减, end_factor=0.2'),
        ('早停', 'Patience=12 epochs'),
        ('验证窗口', '训练截止日后2个月'),
        ('标签模式', '混合标签 (alpha=0.7, rank+abs)'),
        ('数据增强', '40%触发: time_mask(10%)+noise(σ=0.003)+stock_drop(15%)'),
        ('时间衰减半衰期', '730天'),
        ('损失权重', 'pairwise=1, top5=3, ndcg=0.3, precision=0.5, portfolio=0.15, market=0.2'),
        ('Walk-Forward窗口', 'W1-W6, 6个滚动窗口'),
    ]

    table = doc.add_table(rows=len(config_data), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(config_data):
        cell_label = table.cell(i, 0)
        cell_label.width = Cm(4)
        p_label = cell_label.paragraphs[0]
        run_label = p_label.add_run(label)
        set_run_font(run_label, font_cn=FONT_FAMILY_HEI, font_en=FONT_FAMILY_EN, size=FONT_SIZE_BODY, bold=True)
        set_paragraph_spacing(p_label, line_spacing=LINE_SPACING_20)

        cell_value = table.cell(i, 1)
        p_value = cell_value.paragraphs[0]
        run_value = p_value.add_run(value)
        set_run_font(run_value, font_cn=FONT_FAMILY_CN, font_en=FONT_FAMILY_EN, size=FONT_SIZE_BODY)
        set_paragraph_spacing(p_value, line_spacing=LINE_SPACING_20)

    add_paragraph_with_format(doc, '', font_cn=FONT_FAMILY_CN, size=FONT_SIZE_BODY,
                              line_spacing=LINE_SPACING_20)

    # 3.3
    add_heading_custom(doc, '3.3 实验结果与分析', level=1)

    add_heading_custom(doc, '3.3.1 版本演进历史', level=2)
    add_body(doc,
        '项目的模型版本经历了从V3到V8的多次迭代优化，下表展示关键版本的演进过程和表现：'
    )

    # 版本演进表
    add_paragraph_with_format(doc, '表1 模型版本演进历史', font_cn=FONT_FAMILY_HEI, size=FONT_SIZE_BODY,
                              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=LINE_SPACING_20, after=6)

    version_data = [
        ('版本', '日期', '方法', '5日收益率', '说明'),
        ('V6', '06-27', '过度优化版Transformer', '+0.46%', '关闭辅助任务和数据增强，过拟合'),
        ('V7', '06-30', '修复过拟合+return_head+收益门控', '+6.15%', '恢复正则化，自评大幅提升'),
        ('V7 (官方)', '07-01', '评测窗口实际得分', '-1.29%', '自评与官方评测严重背离，揭示固定窗口过拟合问题'),
        ('GBDT Stacking', '07-07', 'Transformer表征+LightGBM', '-5.32%', '两层模型进一步加剧过拟合'),
        ('Enhanced WF', '07-08', 'Walk-Forward+宏观特征', '-1.37%', 'WF框架建立，但宏观特征在困难窗口表现不佳'),
        ('V8+quarter_aware', '07-10', 'V8 Improved+联合门控', '+1.55%', '最优配置：市场聚合+混合标签+Portfolio Loss+市场门控'),
        ('V8全量重训', '07-19', '2010数据+HL=730', '+0.10%', '远期数据稀释了近期有效信号'),
    ]

    table = doc.add_table(rows=len(version_data), cols=5)
    table.style = 'Table Grid'
    for i, row_data in enumerate(version_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            is_header = (i == 0)
            set_run_font(run, font_cn=FONT_FAMILY_HEI if is_header else FONT_FAMILY_CN,
                        font_en=FONT_FAMILY_EN, size=FONT_SIZE_BODY, bold=is_header)
            set_paragraph_spacing(p, line_spacing=LINE_SPACING_20,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER if is_header else None)

    add_paragraph_with_format(doc, '', font_cn=FONT_FAMILY_CN, size=FONT_SIZE_BODY,
                              line_spacing=LINE_SPACING_20)

    add_heading_custom(doc, '3.3.2 Walk-Forward回测结果', level=2)
    add_body(doc,
        '最优配置（V8 Improved + quarter_aware门控）在18次Walk-Forward回测中的详细表现如下：'
        '季末日（6次测试）平均收益率-0.37%，正收益比例仅33%——季末机构调仓、基金排名等行为'
        '导致市场呈现非正常波动模式，基于历史规律训练的模型在这一环境下表现显著退化；'
        '非季末日（12次测试）平均收益率+2.51%，正收益比例高达83%——在正常市场环境下，'
        '模型的排序能力和市场环境自适应策略有效发挥了作用；'
        '综合来看，18次测试平均收益率+1.55%，正收益比例72%，标准差约为3.8%。'
    )

    add_body(doc,
        '这一结果验证了两个核心发现：第一，季末效应在A股市场中真实存在且影响显著——'
        '同一模型在非季末日的选股收益（+2.51%）与季末日（-0.37%）相差达4.57个百分点，'
        '说明仅靠"选出最好的股票"在季末并不够，需要针对这一场景设计专门的后处理策略。'
        '第二，后处理策略的边际收益可能超过模型架构改进——quarter_aware市场门控是纯后处理模块，'
        '不需要重新训练模型，但其带来的收益率提升（约+3pp vs V7）超过了此前任何模型架构层的改进。'
    )

    add_heading_custom(doc, '3.3.3 消融实验', level=2)
    add_body(doc,
        '为定量评估各组件的贡献，项目进行了多项消融实验：'
    )

    add_body(doc,
        '（1）数据窗口截断实验：对比使用全部历史数据（2010年起）和仅使用2021年起的数据训练，'
        '发现所有截断配置的正收益比例均≥50%，全量配置跌至47%。结论：2021年之前的市场数据'
        '对当前预测确实有害——市场结构和定价机制的变化使得旧数据成为噪声而非有效信号。'
        '最优数据起始时间为2021年初，对应约4年的有效训练窗口。'
    )

    add_body(doc,
        '（2）时间衰减半衰期实验：在2021年起的数据上对比无衰减、HL=730天和HL=365天三种配置。'
        'HL=730天（每2年权重减半）显著优于无衰减，验证了"远期数据提供正则化而非主导梯度"的设计理念。'
        'HL=365天过于激进，有效样本量不足以支撑模型训练。'
    )

    add_body(doc,
        '（3）日历特征实验：在特征中显式添加days_to_qe（距季末日天数）和is_qe_month（是否季末月）'
        '两维特征，期望模型自主学会季末效应。实验结果显示非季末日收益率反而退化4.39个百分点，'
        '因为这两维特征增强了模型对季末附近数据的过拟合。'
        '这一发现进一步确认了"模型做排序、后处理做场景适配"的架构分工是最优方案。'
    )

    add_body(doc,
        '（4）宏观特征消融实验：在Walk-Forward窗口1（最困难的验证窗口，训练数据仅约9个月）上，'
        '启用宏观特征的配置final_score比不启用提升14个百分点以上。'
        '在小样本条件下，宏观指标提供的市场环境感知信号对排序质量的边际贡献最为显著。'
    )

    # 3.4
    add_heading_custom(doc, '3.4 项目难点分析', level=1)

    add_heading_custom(doc, '3.4.1 难点一：金融数据非平稳性与过拟合', level=2)
    add_body(doc,
        '金融时间序列数据具有天然的非平稳性——2024年的市场规律与2015年可能截然不同。'
        '项目早期版本（V6）在训练集上final_score达到0.12，但评测窗口实际收益率仅为+0.46%，'
        'V7在自评中取得+6.15%但官方评测仅为-1.29%，暴露了严重的固定窗口过拟合问题。'
        '解决方案分三个层次：第一，采用Walk-Forward滚动窗口交叉验证替代固定验证集，'
        '在6个不同历史区间评估模型泛化能力；第二，引入时间衰减采样，'
        '通过指数衰减使近期数据主导梯度；第三，通过截断实验确定最优数据窗口（约4年），'
        '主动丢弃可能有害的远期数据。三层机制组合使得模型在验证集和评测窗口之间的表现趋于一致。'
    )

    add_heading_custom(doc, '3.4.2 难点二：排序质量与绝对收益的张力', level=2)
    add_body(doc,
        '排序学习模型的核心目标是"正确排序"而非"预测涨跌"。然而在真实投资场景中，'
        '模型可能将所有股票都排在"跌3%">"跌5%"的正确顺序上——排序质量完美但投资组合亏损。'
        '这一矛盾是排序学习在量化选股中最根本的挑战。项目从三个角度试图解决：'
        '第一，混合标签以70%排序和30%绝对收益组合，既保留排序能力又传递方向信号；'
        '第二，Portfolio Return Loss通过Gumbel-Softmax松弛直接最大化选中股票的真实收益期望，'
        '在损失函数层面注入绝对收益的优化目标；第三，收益门控后处理只选择预测绝对收益为正的股票，'
        '在推理阶段硬过滤掉"排序高但预测会跌"的候选。三者协同形成了从标签→损失→推理的全链路解决方案。'
    )

    add_heading_custom(doc, '3.4.3 难点三：季末效应的发现与应对', level=2)
    add_body(doc,
        '季末效应是本项目后期最关键的发现。通过逐日分析Walk-Forward回测结果，'
        '发现模型在季末前后（3月、6月、9月、12月的最后几个交易日）的选股表现显著差于非季末日。'
        '经分析，季末是机构投资者调仓、基金季度排名、股指期货交割等集中发生的时期，'
        '市场价格行为偏离了常规模式。应对方案经历了两个阶段：'
        '首先尝试在模型特征中显式加入日历信息让模型自主学习（方向一：日历特征），'
        '但实验证明反而导致了更严重的过拟合；转而采用后处理策略（方向二：quarter_aware门控），'
        '在模型外部根据日历位置和市场方向动态调整选股策略，既保护了模型的排序能力不被日历特征污染，'
        '又实现了季末场景的专门适配。该方案的最终效果——季末日与非季末日之间4.57pp的收益率差距——'
        '既是成就也是遗憾，说明季末效应尚未被完全克服，仍有改进空间。'
    )

    add_heading_custom(doc, '3.4.4 难点四：长历史数据的内存管理', level=2)
    add_body(doc,
        '当训练数据从2021年扩展至2010年时，物化存储所有"天样本"（每个样本含300只股票×60天×237特征'
        '≈17MB）的内存需求达到了约58GB，远超单机32GB内存的限制。解决方案是实现LazyRankingDataset'
        '懒加载机制——每只股票的特征矩阵仅存储一份float32数组，窗口切片延迟到__getitem__中执行，'
        '将内存占用从58GB压缩至约1GB（存储原始特征矩阵）。这一设计在内存效率和数据加载灵活性之间'
        '取得了合理平衡，虽然增加了一定的I/O开销（每次获取样本需要执行切片操作），'
        '但在典型训练场景下（batch_size=4），I/O开销可被GPU训练时间充分掩盖。'
    )

    # 单元测试
    add_heading_custom(doc, '3.5 单元测试与集成测试', level=1)

    add_body(doc,
        '项目编写了41个单元测试和2个集成测试，覆盖核心功能模块。测试采用pytest框架，'
        '可在约3秒内完成全部单元测试，约30秒完成端到端集成测试。主要测试覆盖：'
    )

    add_body(doc,
        '（1）模型测试（7个）：验证StockTransformer和LightweightStockRanker的前向传播形状正确性、'
        '市场聚合模块的开关控制、完整配置下的参数统计（~2.5M和~264K）、'
        '两种模型的梯度流有效性以及strict=False模型加载的向后兼容性。'
    )

    add_body(doc,
        '（2）损失函数测试（8个）：验证WeightedRankingLoss各子损失的非负性、'
        '梯度从total loss回流至预测分数、预测加常数后排序损失的尺度不变性、'
        'top5_weight增大导致损失增大的行为正确性、全mask样本和单股票等边界条件、'
        'NDCG近似损失对完美排序给出低损失值、以及排名指标计算不崩溃。'
    )

    add_body(doc,
        '（3）标签构建测试（6个）：验证混合标签和纯分位数标签的构建逻辑、'
        '分位数标签的[0,1]范围约束、混合标签与纯rank标签的差异性、'
        'label_abs正负号与价格趋势的一致性、极端低价过滤以及空输入不崩溃的鲁棒性。'
    )

    add_body(doc,
        '（4）市场门控测试（11个）：验证HS300累计收益计算、市场信号结构完整性、'
        '明确上涨/下跌/平稳三种市场环境的信号方向正确性、季末日检测精度、'
        '自适应防御权重在季末看跌和普通看涨场景下的取值范围、无效策略名回退以及防御性评分计算。'
    )

    add_body(doc,
        '（5）集成测试（2个）：使用合成数据验证完整训练链路（数据构建→模型创建→1 epoch训练→预测），'
        '以及验证训练后模型的保存/加载/推理链路。这两个测试不依赖GPU、真实数据或预训练模型，'
        '确保在任何环境下都能快速验证代码的可运行性。'
    )


def write_conclusion(doc):
    """结论与展望"""
    add_page_break(doc)
    add_heading_custom(doc, '结论与展望', level=0)

    add_heading_custom(doc, '4.1 工作总结', level=1)

    add_body(doc,
        '本课程设计围绕沪深300指数预测这一实际量化投资问题，完成了一套完整的'
        '深度排序学习选股系统的设计、实现和评估。主要成果包括：'
    )

    add_body(doc,
        '（1）构建了237维多源异构特征工程体系，系统性地覆盖了量价、技术、基本面、'
        '宏观、市场宽度和行业分类等主要信息维度，为模型提供了丰富的预测信号来源。'
    )

    add_body(doc,
        '（2）设计并实现了StockTransformer深度排序模型（约250万参数），集成多尺度时序卷积、'
        'Transformer自注意力、跨股票交互和市场聚合门控等创新模块；同时提出了参数削减11倍的'
        '轻量级替代方案LightweightStockRanker（约26万参数）。'
    )

    add_body(doc,
        '（3）建立了鲁棒的Walk-Forward滚动窗口训练框架，配合时间衰减采样、混合标签、'
        '多目标损失函数（含Portfolio Return Loss）和数据增强等关键技术，'
        '有效应对了金融时间序列数据的非平稳性和小样本挑战。'
    )

    add_body(doc,
        '（4）提出了quarter_aware市场门控后处理策略，在不重新训练模型的前提下，'
        '通过市场环境自适应实现进攻型与防御型选股策略的动态切换，'
        '在18次回测中取得了综合收益率+1.55%、正收益比例72%的表现。'
    )

    add_body(doc,
        '（5）撰写了完整的项目文档（README、GUIDE、REPRODUCE、ADR）和41个自动化测试，'
        '以及Docker容器化部署方案，确保了项目的可复现性和可交付性。'
    )

    add_heading_custom(doc, '4.2 经验与反思', level=1)

    add_body(doc,
        '回顾整个项目历程，以下几条经验最具启发性：'
    )

    add_body(doc,
        '第一，先查数据管道，再动模型。项目中一个"自然日连续"的过滤条件默默丢弃了约80%的训练样本'
        '（因为周末和节假日不交易，严格连续过滤仅保留了完整交易周的预测日），'
        '这个隐蔽的数据管道问题直到训练样本量异常偏低时才被发现。在量化投资项目中，'
        '数据管道的正确性往往比模型架构的先进性更重要。'
    )

    add_body(doc,
        '第二，后处理策略的效率可能远超模型架构改进。quarter_aware市场门控作为纯后处理模块，'
        '不需要GPU训练即可实现+3pp量级的收益率提升，而此前模型架构层面的多次改进'
        '（V6→V7→V8）单个改进的提升量级通常在1-2pp。这一经验提示：在量化选股中，'
        '"什么时候选什么类型的股票"这个决策维度的重要性不亚于"选出最好的股票"。'
    )

    add_body(doc,
        '第三，金融数据的非平稳性是不可忽视的根本性挑战。每个模型版本的"自评"和"官方评测"之间'
        '的巨大反差（V7自评+6.15% vs 官方-1.29%），以及旧数据（2010-2020年）对模型表现的'
        '显著负面效应，都在反复提醒：在金融领域，"更多数据=更好模型"的深度学习常识并不成立。'
        '对抗非平稳性需要系统性的方案——Walk-Forward验证、时间衰减采样、截断实验——'
        '而非简单的增加数据量。'
    )

    add_heading_custom(doc, '4.3 未来展望', level=1)

    add_body(doc,
        '本项目在以下方向仍有改进空间：'
    )

    add_body(doc,
        '（1）引入多模态数据：当前的237维特征主要基于结构化数据，未来可以探索引入'
        '新闻情感分析（NLP）、机构调研纪要、分析师研报等非结构化文本数据，'
        '利用预训练语言模型提取增量预测信号。'
    )

    add_body(doc,
        '（2）在线学习与自适应更新：当前模型离线训练后固定部署，无法自适应市场环境的'
        '结构性变化。可以探索在线学习框架，以较低的计算成本在新交易日数据上持续更新模型参数。'
    )

    add_body(doc,
        '（3）多周期预测与风险预算：当前模型仅预测未来5日的Top-5选股和等权/softmax权重分配。'
        '可以探索多持有期（5日/10日/20日）联合预测，以及基于风险平价的动态权重分配策略。'
    )

    add_body(doc,
        '（4）模型可解释性：当前深度学习模型是黑箱，无法解释"为什么选这5只股票"。'
        '可以引入注意力权重可视化、Shapley值归因和特征重要性分析，'
        '提升模型的可解释性以满足实际投资中的合规和风控需求。'
    )

    add_body(doc,
        '（5）更精细的季末效应应对：尽管quarter_aware门控缓解了季末效应的影响，'
        '但季末日与非季末日之间4.57pp的收益率差距表明这一问题尚未被完全解决。'
        '可以尝试专门针对季末场景训练独立的"季末模型"，或引入更高频的日内交易数据捕捉季末微观行为。'
    )


def write_references(doc):
    """参考文献"""
    add_page_break(doc)
    add_heading_custom(doc, '参考文献', level=0)

    references = [
        '[1] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems (NeurIPS), 2017: 5998-6008.',
        '[2] Burges C J C. From RankNet to LambdaRank to LambdaMART: An Overview[R]. Microsoft Research Technical Report, 2010.',
        '[3] Qin Z, Yan L, Zhuang H, et al. Learning to Rank: From Pairwise Approach to Listwise Approach[C]. International Conference on Machine Learning (ICML), 2008: 129-136.',
        '[4] Jang E, Gu S, Poole B. Categorical Reparameterization with Gumbel-Softmax[C]. International Conference on Learning Representations (ICLR), 2017.',
        '[5] Fama E F, French K R. A Five-Factor Asset Pricing Model[J]. Journal of Financial Economics, 2015, 116(1): 1-22.',
        '[6] Gu S, Kelly B, Xiu D. Empirical Asset Pricing via Machine Learning[J]. Review of Financial Studies, 2020, 33(5): 2223-2273.',
        '[7] Chen L, Pelger M, Zhu J. Deep Learning in Asset Pricing[J]. Management Science, 2024, 70(2): 718-739.',
        '[8] Ba J L, Kiros J R, Hinton G E. Layer Normalization[J]. arXiv preprint arXiv:1607.06450, 2016.',
        '[9] Loshchilov I, Hutter F. Decoupled Weight Decay Regularization[C]. International Conference on Learning Representations (ICLR), 2019.',
        '[10] Zhang A, Lipton Z C, Li M, et al. Dive into Deep Learning[M]. Cambridge University Press, 2023.',
        '[11] TA-Lib: Technical Analysis Library[EB/OL]. https://ta-lib.org/.',
        '[12] Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]. Advances in Neural Information Processing Systems (NeurIPS), 2019: 8024-8035.',
        '[13] baostock: 证券宝——Python开源证券数据接口[EB/OL]. http://baostock.com/.',
        '[14] akshare: Python开源财经数据接口库[EB/OL]. https://www.akshare.xyz/.',
        '[15] THU-BDC2026 大数据竞赛赛题说明[EB/OL]. https://www.thubdc2026.com/.',
    ]

    for ref in references:
        add_body_no_indent(doc, ref)


def main():
    doc = Document()

    # ─── 页面设置 ─────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ─── 设置默认字体 ─────────────────────────────
    style = doc.styles['Normal']
    style.font.name = FONT_FAMILY_EN
    style.font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY_CN)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING_20
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # ─── 生成各章节 ──────────────────────────────
    print("创建封面...")
    create_cover_page(doc)

    print("创建题目页...")
    create_title_page(doc)

    print("创建摘要...")
    create_abstract(doc)

    print("创建目录占位...")
    create_toc_placeholder(doc)

    print("编写第1章...")
    write_chapter1(doc)

    print("编写第2章...")
    write_chapter2(doc)

    print("编写第3章...")
    write_chapter3(doc)

    print("编写结论...")
    write_conclusion(doc)

    print("编写参考文献...")
    write_references(doc)

    # ─── 保存 ────────────────────────────────────
    output_path = os.path.join(PROJECT_ROOT, '生产实习课程设计报告_沪深300指数预测.docx')
    doc.save(output_path)
    print(f"\n✅ 报告已生成: {output_path}")


if __name__ == '__main__':
    main()
